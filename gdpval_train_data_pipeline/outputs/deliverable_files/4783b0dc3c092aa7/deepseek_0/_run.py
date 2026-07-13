
from docx import Document

doc = Document('Comparative_Inflation_Assessment.docx')

print("=== DOCUMENT STRUCTURE ===")
for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading'):
        print(f"  [{para.style.name}] {para.text}")
    elif i < 15:
        print(f"  [{para.style.name}] {para.text[:100]}...")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

# Check sections present
full_text = "\n".join([p.text for p in doc.paragraphs])
required_sections = [
    "Executive Summary",
    "U.S. Inflation", 
    "Euro Area Inflation",
    "Labor Market and Wages",
    "Risk Assessment"
]

print("\n=== REQUIRED SECTIONS CHECK ===")
for section in required_sections:
    if section in full_text:
        print(f"  [OK] '{section}' found")
    else:
        print(f"  [MISSING] '{section}' NOT found")

# Check key data points
key_data = [
    ("2.1 percent", "Headline PCE April 2025"),
    ("2.5 percent", "Core PCE April 2025"),
    ("2.6 percent", "Headline PCE December 2024"),
    ("2.9 percent", "Core PCE December 2024"),
    ("4.2 percent", "Unemployment rate May 2025"),
    ("4.1 percent", "Unemployment rate December 2024"),
    ("1.9 percent", "PCE food April 2025"),
    ("4.1 percent", "Michigan 5-10yr June 2025"),
    ("5.1 percent", "Michigan 1-yr June 2025"),
]

print("\n=== KEY DATA POINTS CHECK ===")
for value, desc in key_data:
    if value in full_text:
        print(f"  [OK] {desc}: {value}")
    else:
        print(f"  [MISSING] {desc}: {value}")

print("\n=== WORD COUNT ===")
words = full_text.split()
print(f"  Approximate word count: {len(words)}")
