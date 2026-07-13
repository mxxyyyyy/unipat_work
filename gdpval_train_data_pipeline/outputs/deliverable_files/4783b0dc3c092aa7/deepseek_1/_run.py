
import os

path = 'Comparative_Inflation_Assessment.docx'
size = os.path.getsize(path)
print(f'File: {path}')
print(f'Size: {size:,} bytes')

# Quick validation: open and check section count
from docx import Document
doc = Document(path)

print(f'\nParagraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')

# Print headings to verify structure
print('\n--- Document Structure ---')
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        print(f'  [{p.style.name}] {p.text}')
    elif p.text.strip() and len(p.text) > 50:
        print(f'  [Body start] {p.text[:80]}...')

print('\nVerification complete.')
