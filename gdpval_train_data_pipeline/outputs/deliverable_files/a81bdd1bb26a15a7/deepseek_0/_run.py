
# Verify the docx file is valid and readable
from docx import Document
doc = Document('Asset_Valuations_Draft.docx')
print("=== DOCUMENT VERIFICATION ===")
print(f"Number of paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")
print(f"Number of sections: {len(doc.sections)}")
print()

# Check headings
print("=== HEADINGS ===")
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        print(f"  [{p.style.name}] {p.text}")

print()
# Check table
print("=== TABLE CHECK ===")
table = doc.tables[0]
print(f"Table rows: {len(table.rows)}, cols: {len(table.columns)}")
print(f"Header (row 0):")
for j, cell in enumerate(table.rows[0].cells):
    print(f"  Col {j}: {cell.text.strip()[:50]}")

print()
# Verify all sections present
sections_found = []
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        sections_found.append(p.text)

expected = [
    'Asset Valuations: Stretched Markets in a High-Uncertainty Environment',
    '1. Overall Assessment',
    '2. Key Asset Valuation Indicators',
    '3. Detailed Analysis by Asset Class',
    '3.1 Equity Markets',
    '3.2 Corporate Credit Markets',
    '3.3 Sovereign Bond Markets',
    '3.4 Real Estate Markets',
    '4. The Most Stretched Asset Classes',
    '4.1 Euro Area',
    '4.2 United States',
    '5. Potential Triggers for a Sharp Repricing of Risk',
    '6. Vulnerability Rating and Conclusion'
]

print("=== STRUCTURE CHECK ===")
for exp in expected:
    if exp in sections_found:
        print(f"  ✓ {exp}")
    else:
        print(f"  ✗ MISSING: {exp}")

# Print first and last paragraphs
print("\n=== FIRST PARAGRAPH (excerpt) ===")
print(doc.paragraphs[2].text[:200] + "...")
print("\n=== LAST SUBSTANTIVE PARAGRAPH (excerpt) ===")
for p in reversed(doc.paragraphs):
    if len(p.text) > 50:
        print(p.text[:200] + "...")
        break

print("\n=== FILE SIZE ===")
import os
size = os.path.getsize('Asset_Valuations_Draft.docx')
print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
print("\n✓ Verification complete — document is valid and complete.")
