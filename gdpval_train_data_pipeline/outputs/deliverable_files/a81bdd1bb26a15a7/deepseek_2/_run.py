
from docx import Document
doc = Document('Asset_Valuations_Draft.docx')
print("✓ File opens correctly")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables: {len(doc.tables)}")
print(f"  Table rows: {len(doc.tables[0].rows)} x {len(doc.tables[0].columns)}")

# List structure
for para in doc.paragraphs:
    if para.style.name.startswith('Heading'):
        level = para.style.name.split()[-1]
        indent = "  " * (int(level) - 1)
        print(f"{indent}[H{level}] {para.text[:80]}")
