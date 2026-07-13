
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_h(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def scell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = 'Calibri'; r.bold = bold
    if color: r.font.color.rgb = color

# === TITLE ===
title = doc.add_heading('Asset Valuations: Stretched Markets in a High-Uncertainty Environment', level=0)
for run in title.runs: run.font.color.rgb = RGBColor(0, 51, 102)
p = doc.add_paragraph()
run = p.add_run('European Central Bank — Financial Stability Review, November 2025')
run.font.size = Pt(10); run.font.color.rgb = RGBColor(100,100,100); run.italic = True
doc.add_paragraph()

# === 1. OVERVIEW ===
add_h(doc, '1. Overview', 1)
doc.add_paragraph(
    'Euro area asset valuations remain stretched across key markets as of late 2025, even as short-term '
    'volatility has receded from the exceptional levels seen during the April sell-off. The ECB Financial '
    'Stability Review (FSR) of May 2025 warned that "stretched valuations and declining non-bank liquidity '
    'make markets prone to further outsized reactions." While the EURO STOXX 50 and S&P 500 "mostly '
    'recovered their initial losses by mid-May" (ECB FSR, May 2025), the underlying drivers of repricing '
    'risk—elevated trade policy uncertainty, fiscal sustainability concerns, and geopolitical '
    'fragmentation—have intensified. The U.S. Federal Reserve\'s November 2025 FSR corroborates this '
    'picture: the S&P 500 forward P/E ratio is "close to the upper end of its historical range," and '
    'corporate bond spreads have settled "at historically tight levels" (Fed FSR, November 2025). The '
    'compression of risk premia, against deteriorating macro-financial fundamentals, suggests an '
    'uncomfortably narrow margin of safety. This chapter assesses valuation stretch across major asset '
    'classes, identifies the most vulnerable segments, and discusses triggers and amplification mechanisms '
    'for a potential disorderly repricing.'
)

# === 2. TABLE ===
add_h(doc, '2. Key Asset Valuation Indicators', 1)
doc.add_paragraph(
    'Table 1 summarises valuation metrics across equity, fixed-income, real estate, and sovereign '
    'markets, drawing on the ECB FSR (May and November 2025) and the Fed FSR (November 2025).'
)

table_data = [
    ['Asset Class', 'Indicator', 'Region', 'Most Recent Assessment', 'Historical Context', 'Source'],
    ['Equities', 'Forward P/E ratio', 'US (S&P 500)', 'Close to upper end of historical range', '~90th percentile', 'Fed FSR Nov. 2025, Fig. 1.4'],
    ['Equities', 'Equity risk premium (earnings yield – real 10y)', 'US', 'Near a 20-year low', 'Multi-decade trough', 'Fed FSR Nov. 2025, Fig. 1.5'],
    ['Equities', 'EURO STOXX 50 valuation', 'Euro area', 'Stretched; rebounded near pre-sell-off levels', 'High by historical standards', 'ECB FSR May 2025, Overview'],
    ['Equities', 'Option-implied volatility (VIX)', 'US', 'Below historical median', 'Below median; declined from April spike', 'Fed FSR Nov. 2025, Fig. 1.6'],
    ['Corporate Credit', 'IG spreads (triple-B)', 'US', '~0.7 pp below historical median', 'Below median; tight', 'Fed FSR Nov. 2025, Fig. 1.8'],
    ['Corporate Credit', 'HY spreads', 'US', '~1.6 pp below historical median', 'Well below median', 'Fed FSR Nov. 2025, Fig. 1.8'],
    ['Corporate Credit', 'Excess bond premium (EBP)', 'US', 'Below median of historical distribution', 'Below long-run average', 'Fed FSR Nov. 2025, Fig. 1.9'],
    ['Corporate Credit', 'Leveraged loan spreads', 'US', 'Low end of distribution since 2009', 'Bottom of post-2009 range', 'Fed FSR Nov. 2025, Fig. 1.10'],
    ['Corporate Credit', 'Euro area corp. vulnerabilities & credit', 'Euro area', 'High valuations; elevated bankruptcies & DSR', 'Stretched; multi-year highs in bankruptcies', 'ECB FSR May, Nov. 2025'],
    ['Sovereign Bonds', '10y spreads vs. German Bund', 'Euro area', 'Widening for vulnerable issuers; downgrades', 'Distribution widening; fragmentation risk', 'ECB FSR Nov. 2025, Chart 3'],
    ['Sovereign Bonds', 'Govt. debt & projected fiscal balance', 'Euro area', 'Several >100% debt/GDP; deficits near/above 3%', 'Elevated vs. Maastricht thresholds', 'ECB FSR Nov. 2025, Chart 3c'],
    ['Sovereign Bonds', 'Nominal Treasury term premium', 'US', 'Near historical median', 'At median; declined slightly', 'Fed FSR Nov. 2025, Fig. 1.2'],
    ['Real Estate', 'CRE price growth (deflated)', 'US', '–5.6% y/y (to Q2 2025)', 'Below long-run avg. (+5.4%)', 'Fed FSR Nov. 2025, Table 1.1'],
    ['Real Estate', 'Residential real estate price growth', 'US', '+1.4% y/y; nominal +0.3–1.7%', 'Below long-run avg. (+6.2%)', 'Fed FSR Nov. 2025, Table 1.1'],
    ['Real Estate', 'Euro area real estate markets', 'Euro area', 'Recovering but facing headwinds', 'Early recovery; fragile', 'ECB FSR May 2025, Sec. 1.5'],
]

table = doc.add_table(rows=len(table_data), cols=6)
table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = True
for i, h in enumerate(table_data[0]):
    scell(table.rows[0].cells[i], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(255,255,255))
    shade(table.rows[0].cells[i], '003366')
for r in range(1, len(table_data)):
    for c in range(6): scell(table.rows[r].cells[c], table_data[r][c], bold=(c==0), size=8)
    if r % 2 == 0:
        for c in range(6): shade(table.rows[r].cells[c], 'EBF0F7')

p = doc.add_paragraph()
run = p.add_run('Table 1: '); run.bold = True; run.font.size = Pt(9)
run = p.add_run('Key asset valuation indicators — euro area and United States, late 2025')
run.font.size = Pt(9); run.italic = True
p = doc.add_paragraph()
run = p.add_run('Sources: '); run.bold = True; run.font.size = Pt(8)
run = p.add_run('ECB FSR (May 2025, November 2025); Fed FSR (November 2025). "pp" = percentage points. DSR = debt service ratio.')
run.font.size = Pt(8); run.font.color.rgb = RGBColor(80,80,80)

doc.add_paragraph()

# === 3. ASSESSMENT BY ASSET CLASS ===
add_h(doc, '3. Assessment by Asset Class', 1)

# 3.1 Equities
add_h(doc, '3.1 Equity Markets', 2)
doc.add_paragraph(
    'Equity valuations present the clearest signal of stretch. In the United States, the S&P 500 forward '
    'P/E ratio stands "close to the upper end of its historical range"—roughly the 90th percentile—while '
    'the equity risk premium (expected earnings yield minus the real 10-year Treasury yield) is "near a '
    '20-year low" (Fed FSR, November 2025, Figures 1.4–1.5). Investors are accepting unusually thin '
    'compensation for bearing equity risk, a configuration historically associated with elevated '
    'vulnerability to negative shocks.'
)
doc.add_paragraph(
    'Euro area equities exhibit a similarly stretched profile. The EURO STOXX 50 "mostly recovered" '
    'its April losses by mid-May, accompanied by "signs of a rotation in fund flow dynamics from the '
    'United States to the euro area" (ECB FSR, May 2025). While partly reflecting confidence linked to '
    'prospective fiscal expansion, this implies euro area equities have not repriced for the deteriorating '
    'macro-financial outlook. As the ECB FSR (May 2025) cautioned, "the medium-term implications of '
    'tariffs might still challenge risky asset valuations." The rapid post-April compression of implied '
    'volatility—the VIX is now "below the historical median" (Fed FSR, November 2025, Figure 1.6)—masks '
    'the potential for swift re-emergence of stress given persistently elevated policy uncertainty.'
)

# 3.2 Corporate Credit
add_h(doc, '3.2 Corporate Credit Markets', 2)
doc.add_paragraph(
    'Corporate bond markets exhibit the most pronounced disconnect between valuations and fundamentals. '
    'U.S. triple-B spreads are approximately 0.7 percentage points below their historical median; high-yield '
    'spreads roughly 1.6 percentage points below theirs; the excess bond premium is "below the median of '
    'its historical distribution"; and leveraged loan spreads sit at "the low end of their historical '
    'distribution since 2009" (Fed FSR, November 2025, Figures 1.8–1.10). These readings imply that credit '
    'investors are pricing an exceptionally benign default environment.'
)
doc.add_paragraph(
    'This is difficult to reconcile with deteriorating corporate fundamentals. The ECB FSR (November 2025) '
    'notes that "corporate vulnerabilities remain elevated as the effects of tariffs take hold," with '
    'rising bankruptcies, elevated debt service costs, and a corporate vulnerability index in above-average '
    'territory (Charts 1.5–1.6). Manufacturing firms—disproportionately exposed to trade disruptions—display '
    'the weakest Altman Z-scores. The ECB FSR (May 2025) highlighted "high valuations and increasing risk '
    'concentration" in credit markets, a warning that remains pertinent as spreads offer scant buffer '
    'against potential downgrades and defaults.'
)

# 3.3 Sovereign Bonds
add_h(doc, '3.3 Sovereign Bond Markets', 2)
doc.add_paragraph(
    'Euro area sovereign bond markets present pockets of significant vulnerability. Ten-year spreads '
    'versus German Bunds have widened for several member states, with some experiencing credit rating '
    'downgrades since November 2024 (ECB FSR, November 2025, Chart 3). Several countries combine '
    'government debt above 100% of GDP with projected fiscal balances near or exceeding the 3% Maastricht '
    'threshold for 2025–26 (Chart 3c). The ECB FSR (May 2025) noted that "increased defence spending '
    'needs could exacerbate fiscal pressures for some highly indebted sovereigns with substantial short-term '
    'refinancing requirements" (Chart 5); the activation of the national escape clause for 11 euro area '
    'countries adds bond supply pressures. The November 2025 edition cautions that "higher issuance and '
    'funding costs could strain weak sovereigns, with fiscal slippage or external fiscal stress potentially '
    'leading to renewed sovereign fragmentation." In the United States, the nominal Treasury term premium '
    'remains near its historical median (Fed FSR, November 2025, Figure 1.2), though this could shift '
    'rapidly if fiscal sustainability concerns intensify.'
)

# 3.4 Real Estate
add_h(doc, '3.4 Real Estate Markets', 2)
doc.add_paragraph(
    'Real estate markets are in a correction or early recovery phase rather than exhibiting stretched '
    'valuations. U.S. CRE prices declined 5.6% year-on-year (to Q2 2025), below the long-run average '
    'of 5.4%, while residential price growth moderated to 1.4% (Fed FSR, November 2025, Table 1.1). '
    'Euro area real estate markets are "recovering" but "may face headwinds from elevated uncertainty" '
    '(ECB FSR, May 2025, Section 1.5). While not overvalued, the sector\'s interest-rate sensitivity '
    'and its importance on bank and non-bank balance sheets warrant continued monitoring.'
)

# === 4. MOST STRETCHED ===
add_h(doc, '4. Most Stretched Asset Classes', 1)

add_h(doc, '4.1 Euro Area: Equities and Corporate Credit', 2)
doc.add_paragraph(
    'The two most stretched asset classes in the euro area are equities and corporate credit. Equity '
    'valuations have rebounded to near pre-sell-off levels despite a marked growth deterioration—the ECB '
    'FSR (November 2025) notes that "global uncertainties and tariff rates weigh on euro area growth" '
    '(Chart 1.1). The May 2025 FSR characterised valuations as "stretched" and markets as "prone to '
    'further outsized reactions." Fund flow rotation into euro area equities may have amplified the '
    'price-fundamentals divergence. Corporate credit is the second area of acute stretch: "high valuations '
    'and increasing risk concentration" coexist with "elevated" vulnerabilities, rising bankruptcies, and '
    'high debt service ratios (ECB FSR, November 2025). The manufacturing sector, accounting for a '
    'significant share of euro area employment and bank credit, is disproportionately exposed to tariff '
    'shocks (Chart 4). Compressed credit risk premia in this environment represent a material mispricing '
    'of downside risks.'
)

add_h(doc, '4.2 United States: Equities and Corporate Bonds', 2)
doc.add_paragraph(
    'In the United States, equities and corporate bonds are the two most stretched classes, with '
    'overvaluation more precisely quantified. The S&P 500 forward P/E at roughly the 90th percentile '
    'and the equity risk premium at a 20-year low provide unusually clear signals (Fed FSR, November '
    '2025, Figures 1.4–1.5). On the credit side, IG and HY spreads at 0.7 and 1.6 percentage points '
    'below their historical medians, the excess bond premium below its long-run average, and leveraged '
    'loan spreads at the bottom of their post-2009 distribution constitute one of the tightest credit '
    'valuation configurations of the post-GFC period (Figures 1.8–1.10).'
)

# === 5. TRIGGERS ===
add_h(doc, '5. Potential Triggers and Amplification Mechanisms', 1)
doc.add_paragraph(
    'Several catalysts could trigger a sharp repricing across stretched asset classes, interacting '
    'with amplification mechanisms that raise the risk of disorderly propagation.'
)
doc.add_paragraph(
    'Trade policy escalation remains the most immediate trigger. The U.S. effective tariff rate rose '
    'sharply in 2025, while the global trade policy uncertainty index stands far above its 2015–24 '
    'average (ECB FSR, November 2025, Chart 1.1). The April 2025 tariff announcement triggered "a major '
    'sell-off in riskier assets, with magnitudes not seen since the early stages of the COVID-19 pandemic" '
    '(ECB FSR, May 2025). Although markets rebounded after the 90-day tariff pause, the trade policy '
    'trajectory remains unresolved; renewed escalation targeting European exports would directly impact '
    'the manufacturing-intensive sectors where corporate vulnerabilities are already elevated.'
)
doc.add_paragraph(
    'Fiscal sustainability concerns represent a second trigger, particularly for the euro area. '
    'Defence spending needs, high debt levels, and substantial near-term refinancing requirements '
    '"could exacerbate fiscal pressures" (ECB FSR, May 2025, Chart 5). The November 2025 edition warns '
    'of "a potential reassessment of risk by investors" that could reignite the sovereign-bank nexus—'
    'the feedback loop in which deteriorating sovereign creditworthiness impairs bank balance sheets '
    'via bond holdings, and bank distress feeds back to the sovereign through contingent liabilities. '
    'Strengthening funding ties between banks and non-banks create additional contagion channels '
    '(ECB FSR, November 2025, Special Feature on systemic risks in bank-NBFI linkages).'
)
doc.add_paragraph(
    'Third, the NBFI sector could amplify rather than absorb shocks. The ECB FSR (May 2025) documented '
    'increasing leverage among euro area investment funds alongside declining liquidity buffers. Forced '
    'deleveraging via fire sales to meet margin calls or redemptions—consistent with the April 2025 '
    'episode where "an unwinding of asset swap or basis trade positions contribut[ed] to the sell-off '
    'in US Treasuries" (ECB FSR, May 2025)—could transmit stress across markets. The Fed FSR (November '
    '2025) similarly notes that hedge fund leverage "has steadily increased across a broad range of '
    'strategies."'
)
doc.add_paragraph(
    'Fourth, corporate-to-household feedback loops could render sectoral stress systemic. The ECB FSR '
    '(November 2025) observes that corporate vulnerabilities "potentially weaken[...] household resilience '
    'if layoffs affect their debt servicing capacity" (Chart 4). The misery index has risen, and a rolling '
    'recession dynamic (Box 1) could cascade sector-specific downturns into a broader macro-financial shock.'
)

# === 6. VULNERABILITY RATING ===
add_h(doc, '6. Overall Vulnerability Assessment', 1)
doc.add_paragraph(
    'The evidence reviewed supports an overall asset valuation vulnerability rating of ELEVATED for '
    'both the euro area and the United States. This assessment rests on four pillars.'
)
doc.add_paragraph(
    'First, equity valuations are stretched to levels historically associated with heightened '
    'vulnerability—the S&P 500 forward P/E at roughly the 90th percentile and the equity risk premium '
    'at a 20-year low for the United States, with ECB FSR assessments pointing to comparable stretch '
    'in euro area equities.'
)
doc.add_paragraph(
    'Second, credit markets exhibit a pronounced disconnect between historically tight spreads and '
    'deteriorating corporate fundamentals. U.S. IG and HY spreads at 0.7 and 1.6 percentage points below '
    'their historical medians imply a benign default outlook inconsistent with rising bankruptcies, '
    'elevated debt service costs, and intensifying tariff headwinds in the euro area.'
)
doc.add_paragraph(
    'Third, euro area sovereign bond markets face renewed fragmentation risk from high debt levels, '
    'increased defence spending needs, and potential investor reassessment of sovereign credit risk. '
    'The sovereign-bank-NBFI nexus creates channels for rapid cross-asset and cross-border transmission.'
)
doc.add_paragraph(
    'Fourth, stretched valuations coexist with elevated macro-financial uncertainty—trade policy, '
    'geopolitics, and fiscal sustainability—creating an environment where multiple triggers could '
    'materialise simultaneously. The April 2025 episode, when a single policy announcement generated '
    'cross-asset volatility unseen since the COVID-19 shock, illustrates the fragility of the current '
    'equilibrium. Narrow risk premia across asset classes suggest markets are poorly positioned for the '
    'materialisation of the tail risks identified in this Review.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Vulnerability rating: ELEVATED ')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor(180, 50, 0)
run = p.add_run('(scale: Low / Moderate / Elevated / High). This rating reflects the simultaneous '
                'compression of risk premia across equities and corporate credit, set against deteriorating '
                'macro-financial fundamentals and multiple plausible repricing triggers.')
run.font.size = Pt(10); run.italic = True

doc.save('Asset_Valuations_Draft.docx')

# Final word counts
para_words = sum(len(p.text.split()) for p in doc.paragraphs)
heading_words = sum(len(p.text.split()) for p in doc.paragraphs if p.style.name.startswith('Heading'))
body_words = para_words - heading_words
table_words = sum(len(cell.text.split()) for table in doc.tables for row in table.rows for cell in row.cells)
print(f"Paragraph words: {para_words}")
print(f"  Headings: {heading_words}")
print(f"  Body + captions: {body_words}")
print(f"Table words: {table_words}")
print(f"Grand total: {para_words + table_words}")
print(f"\nEstimated prose-only (body minus captions): ~{body_words - 35} words")

# Verify
doc2 = Document('Asset_Valuations_Draft.docx')
print(f"\n✓ File verified: {len(doc2.paragraphs)} paragraphs, {len(doc2.tables)} table, openable")
